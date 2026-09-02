# -*- coding: utf-8 -*-
"""
主生成脚本 — 读取文档数据，生成 PDF / Word / Markdown 三种格式
在 llmdev 环境运行: conda run -n llmdev python gen_all.py
"""
import os, sys
from pathlib import Path

# 注册中文字体
try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for nm, p in [("SimSun","C:/Windows/Fonts/simsun.ttc"),("SimHei","C:/Windows/Fonts/simhei.ttf")]:
        try: pdfmetrics.registerFont(TTFont(nm, p))
        except: pass
except ImportError:
    print("reportlab not found, PDF generation will be skipped")

# 导入文档数据
sys.path.insert(0, str(Path(__file__).parent))
from doc_data_1 import DOCS as D1
from doc_data_2 import DOCS as D2
from doc_data_3 import DOCS as D3

ALL_DOCS = D1 + D2 + D3

OUT_DIR = Path(__file__).parent / "generated"
OUT_DIR.mkdir(exist_ok=True)


def gen_markdown(doc, path):
    """生成 Markdown 文件"""
    lines = [f"# {doc['title']}\n"]
    lines.append(f"> 类别: {doc['category']} | 作者: {doc['author']} | 日期: {doc['date']}\n")
    for sec in doc['sections']:
        lines.append(f"\n## {sec['h1']}\n")
        for p in sec['content']:
            lines.append(p + "\n")
        for sub in sec.get('subsections', []):
            lines.append(f"\n### {sub['h2']}\n")
            for p in sub['content']:
                lines.append(p + "\n")
    Path(path).write_text("\n".join(lines), encoding='utf-8')
    print(f"  MD  -> {path}")


def gen_word(doc, path):
    """生成 Word 文件"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    d.styles['Normal'].font.size = Pt(10.5)
    d.styles['Normal'].font.name = 'SimSun'

    # 标题
    t = d.add_heading(doc['title'], level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"类别: {doc['category']}  |  作者: {doc['author']}  |  日期: {doc['date']}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)

    # 正文
    for sec in doc['sections']:
        d.add_heading(sec['h1'], level=1)
        for para in sec['content']:
            d.add_paragraph(para)
        for sub in sec.get('subsections', []):
            d.add_heading(sub['h2'], level=2)
            for para in sub['content']:
                d.add_paragraph(para)

    d.save(str(path))
    print(f"  DOCX-> {path}")


def gen_pdf(doc, path):
    """生成 PDF 文件"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors

    body = ParagraphStyle('body', fontName='SimSun', fontSize=10, leading=16, firstLineIndent=20)
    h1st = ParagraphStyle('h1', fontName='SimHei', fontSize=16, leading=24, spaceBefore=12, spaceAfter=6)
    h2st = ParagraphStyle('h2', fontName='SimHei', fontSize=13, leading=20, spaceBefore=8, spaceAfter=4)
    title_st = ParagraphStyle('title', fontName='SimHei', fontSize=22, leading=30, alignment=1, spaceAfter=16)
    meta_st = ParagraphStyle('meta', fontName='SimSun', fontSize=9, leading=14, alignment=1, textColor=colors.grey, spaceAfter=20)

    story = []
    story.append(Paragraph(doc['title'], title_st))
    story.append(Paragraph(f"类别: {doc['category']} | 作者: {doc['author']} | 日期: {doc['date']}", meta_st))

    for sec in doc['sections']:
        story.append(Paragraph(sec['h1'], h1st))
        for p in sec['content']:
            story.append(Paragraph(p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body))
        for sub in sec.get('subsections', []):
            story.append(Paragraph(sub['h2'], h2st))
            for p in sub['content']:
                story.append(Paragraph(p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body))

    SimpleDocTemplate(
        str(path), pagesize=A4,
        topMargin=2*cm, bottomMargin=2*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm
    ).build(story)
    print(f"  PDF -> {path}")


def main():
    print(f"\n{'='*60}")
    print(f"  生成 {len(ALL_DOCS)} 篇垂直行业文档 (PDF + Word + Markdown)")
    print(f"{'='*60}\n")

    for i, doc in enumerate(ALL_DOCS, 1):
        print(f"[{i}/{len(ALL_DOCS)}] {doc['title']}")
        base = OUT_DIR / doc['filename']
        gen_markdown(doc, str(base) + ".md")
        gen_word(doc, str(base) + ".docx")
        gen_pdf(doc, str(base) + ".pdf")
        print()

    print(f"{'='*60}")
    print(f"  完成! 共生成 {len(ALL_DOCS)*3} 个文件")
    print(f"  输出目录: {OUT_DIR}")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
